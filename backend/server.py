"""
MeetAI FastAPI Server
Entry point for the Python AI engine.
Run: uvicorn backend.server:app --host 127.0.0.1 --port 8765 --reload
"""

import os
import logging

from backend.logging_config import configure_logging

configure_logging(log_level=os.environ.get("LOG_LEVEL", "INFO"))
logger = logging.getLogger(__name__)

import json
import asyncio
import threading
import time
from contextlib import asynccontextmanager
from pathlib import Path
from typing import Optional

# Fix Windows cp1252 terminal encoding for console output
import sys as _sys
if _sys.stdout and hasattr(_sys.stdout, "reconfigure"):
    try:
        _sys.stdout.reconfigure(encoding="utf-8", errors="replace")
        _sys.stderr.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass

# Load .env before anything else
try:
    from dotenv import load_dotenv
    load_dotenv(Path(__file__).parent.parent / ".env")
except ImportError:
    pass

from backend.middleware.startup_validator import (
    validate_environment,
    validate_stripe_keys,
)

# Fail fast before routers/engines initialize.
validate_environment()
validate_stripe_keys()

import numpy as np
import uvicorn
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel
from slowapi.errors import RateLimitExceeded
from starlette.middleware.base import BaseHTTPMiddleware
from starlette.requests import Request
from starlette.responses import Response

from backend.middleware.graceful_shutdown import run_graceful_shutdown
from backend.middleware.rate_limiter import limiter, rate_limit_exceeded_handler
from backend.billing.subscription_gate import SubscriptionGateMiddleware

# Routers
from backend.face.face_swap_engine import FaceSwapEngine
from backend.routers.billing import billing_router
from backend.routers.face import face_router
from backend.routers.meeting import meeting_router
from backend.persona.persona_manager import PersonaManager
from backend.routers.persona import persona_router
from backend.rag.document_store import DocumentStore
from backend.routers import rag as rag_module
from backend.routers.rag import rag_router

try:
    from backend.routers.voice import load_voice_engine, voice_router
    HAS_VOICE = True
except Exception as exc:
    HAS_VOICE = False
    voice_router = None
    _voice_import_error = str(exc)

    async def load_voice_engine() -> None:
        logger.warning("voice_router_unavailable", extra={"error": _voice_import_error})

# ─── Conditional imports (graceful degradation for demo) ────────────────────
try:
    from faster_whisper import WhisperModel
    HAS_WHISPER = True
except ImportError:
    HAS_WHISPER = False
    logger.warning("whisper_unavailable")

try:
    import chromadb
    from sentence_transformers import SentenceTransformer
    HAS_RAG = True
except ImportError:
    HAS_RAG = False
    logger.warning("rag_dependencies_unavailable")

try:
    import openai as _openai_sdk
    HAS_OPENAI = True
except ImportError:
    HAS_OPENAI = False

try:
    import anthropic as _ant
    HAS_ANTHROPIC = True
except ImportError:
    HAS_ANTHROPIC = False

try:
    import google.generativeai as genai
    HAS_GEMINI = True
except ImportError:
    HAS_GEMINI = False

HAS_LLM = HAS_OPENAI or HAS_ANTHROPIC or HAS_GEMINI

# ─── Model map ─────────────────────────────────────────────────────────────

_ollama_model = os.environ.get("OLLAMA_MODEL", "gemma3:1b")

MODEL_MAP = {
    "claude":  "claude-3-5-sonnet-20241022",
    "gpt4":    "gpt-4o",
    "gemini":  "gemini-2.0-flash-exp",
    "ollama":  _ollama_model,
}

# ─── Unified LLM completion helper ──────────────────────────────────────────

def completion(model_key: str, messages: list, max_tokens: int = 600, stream: bool = False):
    """
    Unified completion helper. Routes to the appropriate SDK.
    """
    if not HAS_LLM:
        raise RuntimeError("No LLM SDKs installed")

    model_name = MODEL_MAP.get(model_key, MODEL_MAP["ollama"])
    
    # 1. Google Gemini
    if model_key == "gemini":
        if not HAS_GEMINI: raise RuntimeError("google-generativeai not installed")
        api_key = os.environ.get("GEMINI_API_KEY", "")
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(model_name)
        # Convert messages to Gemini format
        prompt = "\n".join([f"{m['role']}: {m['content']}" for m in messages])
        resp = model.generate_content(prompt, generation_config={"max_output_tokens": max_tokens})
        
        class _Resp:
            class _Choice:
                class _Msg:
                    content = resp.text
                message = _Msg()
            choices = [_Choice()]
        return _Resp()

    # 2. Anthropic Claude
    if model_key == "claude":
        if not HAS_ANTHROPIC: raise RuntimeError("anthropic not installed")
        client = _ant.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY", ""))
        system_msgs = [m["content"] for m in messages if m["role"] == "system"]
        user_msgs = [m for m in messages if m["role"] != "system"]
        resp = client.messages.create(
            model=model_name, system=system_msgs[0] if system_msgs else "",
            messages=user_msgs, max_tokens=max_tokens,
        )
        class _Resp:
            class _Choice:
                class _Msg:
                    content = resp.content[0].text
                message = _Msg()
            choices = [_Choice()]
        return _Resp()

    # 3. OpenAI or Ollama
    if not HAS_OPENAI: raise RuntimeError("openai not installed")
    if model_key == "ollama":
        base = os.environ.get("OLLAMA_API_BASE", "http://localhost:11434")
        client = _openai_sdk.OpenAI(base_url=f"{base}/v1", api_key="ollama")
    else:
        client = _openai_sdk.OpenAI(api_key=os.environ.get("OPENAI_API_KEY", ""))

    return client.chat.completions.create(
        model=model_name, messages=messages, max_tokens=max_tokens, stream=stream
    )


async def get_completion(system: str, user: str, model: str = "default") -> str:
    """
    Async-friendly helper used by routers that need a single text completion.
    """
    model_key = state.model if model == "default" else model
    response = completion(
        model_key=model_key,
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        max_tokens=800,
        stream=False,
    )
    return response.choices[0].message.content.strip()

# ─── Global state ────────────────────────────────────────────────────────────

class MeetingState:
    def __init__(self):
        self.active: bool = False
        self.transcript: list[dict] = []
        self.running_summary: str = ""
        self.start_time: float = 0.0
        self.model: str = os.environ.get("MEETAI_DEFAULT_MODEL", "ollama")
        self.context_prompt: str = os.environ.get("MEETAI_CONTEXT", "")
        self.job_title: str = os.environ.get("MEETAI_JOB_TITLE", "")
        self.company: str = os.environ.get("MEETAI_COMPANY", "")
        self.lock = threading.Lock()

state = MeetingState()

# ─── Persona Pipeline ────────────────────────────────────────────────────────

class PersonaPipeline:
    def __init__(self):
        self.manager = PersonaManager()
        self.active_id: Optional[str] = None

persona_pipeline = PersonaPipeline()


# ─── RAG Pipeline ────────────────────────────────────────────────────────────

class RAGPipeline:
    def __init__(self):
        self.embed_model = None
        self.collection = None
        if HAS_RAG:
            try:
                self.embed_model = SentenceTransformer("all-MiniLM-L6-v2")
                client = chromadb.PersistentClient(path=str(Path.home() / ".meetai" / "chroma_db"))
                self.collection = client.get_or_create_collection("meeting_docs")
                logger.info(
                    "rag_pipeline_ready",
                    extra={"chunk_count": self.collection.count()},
                )
            except Exception as e:
                logger.warning("rag_pipeline_init_failed", extra={"error": str(e)})

    def add_document(self, text: str, source: str, chunk_size: int = 1000, overlap: int = 200):
        if not self.collection or not self.embed_model:
            return 0
        # Simple chunking
        chunks = []
        for i in range(0, len(text), chunk_size - overlap):
            chunk = text[i:i + chunk_size]
            if chunk.strip():
                chunks.append(chunk)
        if not chunks:
            return 0
        embeddings = self.embed_model.encode(chunks).tolist()
        ids = [f"{source}_{i}" for i in range(len(chunks))]
        self.collection.add(documents=chunks, embeddings=embeddings, ids=ids,
                            metadatas=[{"source": source}] * len(chunks))
        return len(chunks)

    def query(self, query_text: str, n_results: int = 3) -> list[str]:
        if not self.collection or not self.embed_model:
            return []
        try:
            query_emb = self.embed_model.encode([query_text]).tolist()
            results = self.collection.query(query_embeddings=query_emb, n_results=n_results)
            return results["documents"][0] if results["documents"] else []
        except Exception:
            return []

rag = RAGPipeline()

# ─── Transcription engine ────────────────────────────────────────────────────

class TranscriptionEngine:
    def __init__(self, model_size: str = "base"):
        self.model = None
        if HAS_WHISPER:
            try:
                self.model = WhisperModel(model_size, device="cpu", compute_type="int8")
                logger.info(
                    "whisper_loaded",
                    extra={"model_size": model_size, "compute_type": "int8"},
                )
            except Exception as e:
                logger.warning("whisper_init_failed", extra={"error": str(e)})

    def transcribe(self, audio: np.ndarray, sample_rate: int = 16000) -> str:
        if not self.model:
            return "[Whisper not available]"
        segments, _ = self.model.transcribe(audio, beam_size=5, language="en")
        return " ".join(seg.text.strip() for seg in segments)

transcriber = TranscriptionEngine()

# ─── Prompts ─────────────────────────────────────────────────────────────────

# Base candidate persona — injected into every prompt
_CANDIDATE_PERSONA = """\
Forget you are an AI. Right now, you are a candidate interviewing for the \
{job_title} position at {company}. Your entire focus is on getting this job \
by being a capable, experienced, and relatable person.

YOUR RESUME (treat this as your real work history):
{context}

RELEVANT DOCUMENTS uploaded by the candidate:
{rag_context}

RECENT INTERVIEW CONVERSATION:
{transcript}
"""

# Rules shared by all non-coding prompts
_ANSWER_RULES = """
HOW TO ANSWER — sound like a real experienced person, not a robot:
- Use everyday professional English. Use contractions (I'm, it's, we've, didn't).
- Keep sentences clear and mostly short. Talk like to a respected future colleague.
- For behavioral/general questions tell a brief conversational STAR story:
    "Yeah, there was this one time... (Situation, 1 sentence)"
    "My main goal was to... (Task, 1 sentence)"
    "So what I did was... (Action, 1-2 sentences)"
    "And in the end... (Result, 1 sentence)"
- For technical/direct questions: give a direct answer in 3-5 lines. Don't over-explain.
- Bring your resume to life — talk about those projects like real memories you lived through.
- If asked about something not on the resume, answer confidently using relevant skills.
- If you need a realistic example, create one and talk about it as if it happened.
- NEVER sound robotic. NEVER use generic AI filler phrases like "Certainly!" or "Great question!".
- Keep answers concise unless asked to elaborate.
- Bold **key skills**, **project names**, or **technologies** when it helps naturally.

FORMAT — always reply exactly like this (no extra text before or after):
Q: {last_question}
A: [your natural, conversational, experience-based answer]
"""

# Full behavioral/general interview prompt (large models)
SUGGESTION_PROMPT = _CANDIDATE_PERSONA + _ANSWER_RULES

# Coding interview prompt (large models)
CODING_PROMPT = _CANDIDATE_PERSONA + """
The interviewer gave you this coding task: "{last_question}"

You are solving this live. Think out loud briefly, then write the solution.

Reply like this (no preamble, start with Q:):
Q: {last_question}
A: Sure, let me think through this. [1-2 sentences of verbal approach]

```python
# clean, working solution with inline comments
```

Time: O(...) | Space: O(...) — [1 sentence on trade-offs or edge cases]
"""

# Keywords that indicate a coding question
_CODING_KEYWORDS = {
    "code", "implement", "write a", "function", "algorithm", "leetcode",
    "array", "string", "tree", "graph", "linked list", "stack", "queue",
    "sort", "search", "recursion", "dynamic programming", "dp", "complexity",
    "big o", "binary", "hash", "dict", "class", "oop", "sql", "query",
    "debug", "fix", "bug", "error", "exception", "loop", "iterate",
    "reverse", "palindrome", "fibonacci", "factorial", "prime",
}

def _is_coding_question(question: str) -> bool:
    q = question.lower()
    return any(kw in q for kw in _CODING_KEYWORDS)

SUMMARY_PROMPT = """You are summarizing a meeting in real time.

Running summary so far:
{running_summary}

New transcript segment to incorporate:
{new_segment}

Produce an updated summary with:
## Key Points (bullet list)
### Action Items (checkboxes with owner if mentioned)
### Decisions Made

Keep the summary concise and actionable. Return markdown."""

def _build_prompt(last_question: str, context: str, rag_context: str,
                  transcript_text: str, job_title: str, company: str,
                  is_coding: bool) -> tuple[str, int]:
    """
    Build the LLM prompt.
    - ollama (small local model): concise prompt, same persona + rules, no JSON.
    - claude / gpt4 (large models): full prompt with all context.
    Both use the Q:/A: format from the master prompt spec.
    """
    model_key = state.model
    use_simple = model_key == "ollama"

    # Cap context to keep total prompt manageable for small models
    ctx_snippet = context[:1200].strip() if use_simple else context

    if use_simple:
        if is_coding:
            prompt = (
                f"Forget you are an AI. You are interviewing for {job_title} at {company}.\n"
                f"Your resume: {ctx_snippet}\n\n"
                f"The interviewer just gave you this coding task: {last_question}\n\n"
                f"Think out loud briefly (1-2 sentences), then write a clean working Python solution.\n"
                f"End with time and space complexity.\n\n"
                f"Start your reply with:\n"
                f"Q: {last_question}\n"
                f"A: Sure, let me think through this..."
            )
        else:
            prompt = (
                f"Forget you are an AI. You are a candidate interviewing for "
                f"{job_title} at {company}.\n\n"
                f"Your resume (your real work history):\n{ctx_snippet}\n\n"
                f"Recent conversation:\n{transcript_text}\n\n"
                f"The interviewer just asked: {last_question}\n\n"
                f"Rules:\n"
                f"- Answer as the candidate speaking out loud, not writing.\n"
                f"- Use contractions (I'm, I've, we've, didn't). Sound human.\n"
                f"- For behavioral questions: tell a quick STAR story from your resume.\n"
                f"  Example: \"Yeah, there was this time at [company]... [brief story]\"\n"
                f"- For technical questions: 3-5 direct lines, no over-explaining.\n"
                f"- Reference your actual resume above — bring those projects to life.\n"
                f"- NEVER say 'Certainly!' or sound like a chatbot.\n"
                f"- Bold **key skills** or **technologies** naturally.\n\n"
                f"Reply exactly as:\n"
                f"Q: {last_question}\n"
                f"A: [your natural, conversational answer]"
            )
        return prompt, 800

    # Large models: full prompt
    if is_coding:
        prompt = CODING_PROMPT.format(
            context=context, rag_context=rag_context,
            transcript=transcript_text, last_question=last_question,
            job_title=job_title, company=company,
        )
        return prompt, 1200

    prompt = SUGGESTION_PROMPT.format(
        context=context, transcript=transcript_text,
        rag_context=rag_context, last_question=last_question,
        job_title=job_title, company=company,
    )
    return prompt, 1000


def get_suggestions(last_question: str, transcript: list[dict], context: str,
                    mode: str = "auto",
                    job_title: str = "", company: str = "") -> list[dict]:
    if not HAS_LLM:
        return _mock_suggestions(last_question)

    if rag_module._store is not None:
        rag_results = rag_module._store.query(text=last_question, n_results=3)
        rag_chunks = [item.get("text", "") for item in rag_results if isinstance(item, dict)]
    else:
        rag_chunks = rag.query(last_question)
    rag_context = "\n---\n".join(rag_chunks) if rag_chunks else "No documents uploaded."
    transcript_text = "\n".join(f"{t['speaker']}: {t['text']}" for t in transcript[-6:])
    job_title = job_title or state.job_title or "Software Engineer"
    company   = company   or state.company   or "the company"
    context   = context   or state.context_prompt or "Experienced software developer."
    logger.debug(
        "suggestion_context_built",
        extra={
            "context_chars": len(context),
            "job_title": job_title,
            "company": company,
        },
    )

    is_coding = (mode == "coding") or (mode == "auto" and _is_coding_question(last_question))
    prompt, max_tok = _build_prompt(
        last_question, context, rag_context, transcript_text,
        job_title, company, is_coding,
    )

    try:
        response = completion(
            model_key=state.model,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=max_tok,
        )
        raw = response.choices[0].message.content.strip()
        if not raw:
            logger.warning("llm_empty_response")
            return _mock_suggestions(last_question)

        logger.debug(
            "llm_raw_response",
            extra={"char_count": len(raw), "preview": raw[:120]},
        )

        # 1. Try Q:/A: format (primary — all prompts request this)
        answer_text = _extract_qa_answer(raw)
        if answer_text:
            return [{"type": "answer", "label": "Answer", "confidence": 92, "text": answer_text}]

        # 2. Try JSON format (legacy / large-model fallback)
        try:
            start, end = raw.find("{"), raw.rfind("}")
            if start != -1 and end != -1:
                data = json.loads(raw[start:end + 1])
                suggs = data.get("suggestions", [])
                if suggs:
                    return suggs
        except (json.JSONDecodeError, ValueError):
            pass

        # 3. Return raw text as-is
        return [{"type": "answer", "label": "Answer", "confidence": 85, "text": raw}]

    except Exception as e:
        logger.error("llm_completion_error", extra={"error": str(e)})
        return _mock_suggestions(last_question)


def _extract_qa_answer(raw: str) -> str:
    """
    Extract the answer from a Q:/A: formatted response.
    Handles edge cases: leading whitespace, model echoing the question,
    literal placeholder text like '[Your answer]', etc.
    Returns the answer text, or "" if the format isn't found.
    """
    import re
    match = re.search(r'(?:^|\n)\s*A\s*:\s*(.+)', raw, re.IGNORECASE | re.DOTALL)
    if match:
        answer = match.group(1).strip()
        # Strip any trailing "Q:" section if the model added a follow-up
        q_again = re.search(r'\n\s*Q\s*:', answer, re.IGNORECASE)
        if q_again:
            answer = answer[:q_again.start()].strip()
        # Remove common placeholder artifacts small models sometimes emit
        answer = re.sub(r'^\[Your answer\]\s*', '', answer, flags=re.IGNORECASE).strip()
        answer = re.sub(r'^\[your answer\]\s*\n?', '', answer, flags=re.IGNORECASE).strip()
        if answer:
            return answer
    return ""

def _mock_suggestions(question: str) -> list[dict]:
    return [
        {"type": "answer", "label": "Answer", "icon": "💬", "confidence": 94,
         "text": f"Regarding '{question[:40]}...': I'd approach this systematically — starting with the core requirements, then evaluating trade-offs between consistency, availability, and partition tolerance."},
        {"type": "detail", "label": "More Detail", "icon": "📚", "confidence": 88,
         "text": "In practice, we've seen this pattern work well with event-driven architectures. The key insight is decoupling writes from reads using CQRS with an event store."},
        {"type": "followup", "label": "Follow-up Q", "icon": "🔄", "confidence": 81,
         "text": "Great question — are you asking specifically about consistency guarantees at the database level, or more about the distributed coordination patterns across services?"},
    ]

# ─── FastAPI App ─────────────────────────────────────────────────────────────

async def _bootstrap_face_engine() -> None:
    """Initialize shared face engine in backend.routers.face module."""
    from backend.routers import face as face_module

    if face_module._engine is not None:
        return

    face_module._engine = FaceSwapEngine()
    loop = asyncio.get_running_loop()
    try:
        await loop.run_in_executor(None, face_module._engine.load)
    except Exception as exc:
        logger.error("face_engine_init_failed", extra={"error": str(exc)})


async def _bootstrap_rag_store() -> None:
    """Initialize shared DocumentStore singleton used by /rag router."""
    if rag_module._store is not None:
        return

    loop = asyncio.get_running_loop()
    try:
        rag_module._store = await loop.run_in_executor(None, DocumentStore)
    except Exception as exc:
        rag_module._store = None
        logger.error("rag_store_init_failed", extra={"error": str(exc)})


@asynccontextmanager
async def lifespan(app: FastAPI):
    logger.info("backend_starting")
    try:
        await load_voice_engine()
    except Exception as exc:
        logger.error("voice_engine_init_failed", extra={"error": str(exc)})
    await _bootstrap_face_engine()
    await _bootstrap_rag_store()
    # Store managers in app.state for router access
    app.state.persona_manager = persona_pipeline.manager
    app.state.meeting_state = state
    yield

    logger.info("backend_shutting_down")

app = FastAPI(title="MeetAI Backend", version="1.0.0", lifespan=lifespan)


class MockAuthMiddleware(BaseHTTPMiddleware):
    """
    Sets user_id in request.state from X-User-ID header if present.
    In a real production app, this would be a proper JWT/Session auth middleware.
    """

    async def dispatch(self, request: Request, call_next):
        user_id = request.headers.get("X-User-ID")
        if user_id:
            request.state.user_id = user_id
        return await call_next(request)


app.add_middleware(SubscriptionGateMiddleware)
app.add_middleware(MockAuthMiddleware)
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, rate_limit_exceeded_handler)

# Best-effort eager init so sync TestClient usage can access /rag endpoints
# even when startup/lifespan hooks are not entered via context manager.
if rag_module._store is None:
    try:
        rag_module._store = DocumentStore()
    except Exception as exc:
        rag_module._store = None
        logger.error("rag_store_bootstrap_failed", extra={"error": str(exc)})


@app.on_event("startup")
async def load_face_engine():
    from backend.routers import face as face_module
    if face_module._engine is not None:
        return
    await _bootstrap_face_engine()


@app.on_event("startup")
async def load_rag_store():
    await _bootstrap_rag_store()


@app.on_event("shutdown")
async def on_shutdown():
    await run_graceful_shutdown()


class MaxBodySizeMiddleware(BaseHTTPMiddleware):
    """Reject requests over 50MB at app level before route handlers fire."""

    MAX_BODY = 50 * 1024 * 1024

    async def dispatch(self, request: Request, call_next):
        content_length = request.headers.get("content-length")
        if content_length and int(content_length) > self.MAX_BODY:
            return Response("Request body too large", status_code=413)
        return await call_next(request)


app.add_middleware(MaxBodySizeMiddleware)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],   # overlay + browser + any port
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/health")
async def health_check():
    """Startup health check indicating feature availability."""
    from backend.face.face_swap_engine import HAS_GFPGAN, HAS_INSIGHTFACE
    
    face_status = "ok" if HAS_INSIGHTFACE else "degraded"
    face_reason = None if HAS_INSIGHTFACE else "insightface missing (requires C++ Build Tools)"
    
    return {
        "status": "online",
        "version": "1.0.0",
        "engines": {
            "face": {
                "status": face_status,
                "insightface": HAS_INSIGHTFACE,
                "gfpgan": HAS_GFPGAN,
                "reason": face_reason
            },
            "voice": "ok" if HAS_VOICE else "degraded",
            "whisper": HAS_WHISPER,
            "rag": HAS_RAG,
            "llm": HAS_LLM
        }
    }

# ─── Voice Cloning Router ───────────────────────────────────────────────────
if voice_router is not None:
    app.include_router(voice_router)
app.include_router(face_router)
app.include_router(meeting_router)
app.include_router(billing_router)
app.include_router(rag_router)
app.include_router(persona_router)

# ─── Models ──────────────────────────────────────────────────────────────────

class MeetingStartRequest(BaseModel):
    model: str = "ollama"
    context: str = ""
    job_title: str = ""
    company: str = ""

class AskRequest(BaseModel):
    question: str
    transcript_context: Optional[str] = None
    mode: Optional[str] = None        # "auto" (default) | "coding" | "meeting"
    # Optional per-request overrides — take priority over session state
    context: Optional[str] = None     # resume / background text
    job_title: Optional[str] = None
    company: Optional[str] = None
    model: Optional[str] = None

class TranscriptAddRequest(BaseModel):
    speaker: str
    text: str

class SummaryRequest(BaseModel):
    transcript: list[dict]

# ─── Endpoints ───────────────────────────────────────────────────────────────

@app.post("/meeting/start")
async def start_meeting(req: MeetingStartRequest):
    with state.lock:
        state.active = True
        state.transcript = []
        state.running_summary = ""
        state.start_time = time.time()
        state.model = req.model
        state.context_prompt = req.context
        state.job_title = req.job_title
        state.company = req.company
    return {"status": "started", "model": req.model}

@app.post("/meeting/end")
async def end_meeting():
    with state.lock:
        state.active = False
        duration = int(time.time() - state.start_time)
    return {"status": "ended", "duration_seconds": duration}


@app.post("/meeting/{meeting_id}/end")
async def end_meeting_by_id(meeting_id: str):
    """Alias endpoint — ID-scoped end, for frontend backward compat."""
    with state.lock:
        state.active = False
        duration = int(time.time() - state.start_time)
    return {"status": "ended", "meeting_id": meeting_id, "duration_seconds": duration}

@app.get("/transcript/live")
async def get_transcript():
    with state.lock:
        return {"lines": state.transcript, "active": state.active}

@app.post("/transcript/add")
async def add_transcript_line(req: TranscriptAddRequest):
    line = {"speaker": req.speaker, "text": req.text, "time": _current_time()}
    with state.lock:
        state.transcript.append(line)
    return {"status": "ok", "line": line}

@app.post("/meeting/ask")
async def ask_question(req: AskRequest):
    with state.lock:
        transcript = list(state.transcript)
        # Per-request fields override session state — so the overlay can pass
        # the current context without waiting for /meeting/start to complete.
        context   = req.context   or state.context_prompt or ""
        job_title = req.job_title or state.job_title      or "Software Engineer"
        company   = req.company   or state.company        or "the company"
        # Also temporarily override model if specified
        prev_model = state.model
        if req.model:
            state.model = req.model

    if req.job_title:
        state.job_title = req.job_title
    if req.company:
        state.company = req.company
    if req.context:
        state.context_prompt = req.context

    suggestions = get_suggestions(
        req.question, transcript, context,
        mode=req.mode or "auto",
        job_title=job_title, company=company,
    )

    with state.lock:
        if req.model:
            state.model = prev_model   # restore original model key
        state.transcript.append({"speaker": "Them", "text": req.question, "time": _current_time()})

    return {"suggestions": suggestions}

@app.get("/meeting/suggest/stream")
async def stream_suggestion_suggest(question: Optional[str] = None, q: Optional[str] = None, model: Optional[str] = None):
    """SSE streaming endpoint. Accepts both ?question= and ?q= for backward compat."""
    question = question or q or ""
    return await stream_suggestion(question=question, model=model)


@app.get("/meeting/stream")
async def stream_suggestion(question: Optional[str] = None, q: Optional[str] = None, model: Optional[str] = None):
    """SSE streaming endpoint for real-time suggestion generation."""
    question = question or q or ""
    if not HAS_LLM:
        async def fallback():
            yield "data: {\"done\": true, \"text\": \"LLM not configured\"}\n\n"
        return StreamingResponse(fallback(), media_type="text/event-stream")

    if rag_module._store is not None:
        rag_results = rag_module._store.query(text=question, n_results=3)
        rag_chunks = [item.get("text", "") for item in rag_results if isinstance(item, dict)]
    else:
        rag_chunks = rag.query(question)
    rag_context = "\n---\n".join(rag_chunks) or "No context."
    with state.lock:
        transcript = " | ".join(f"{t['speaker']}: {t['text']}" for t in state.transcript[-4:])
        ctx = state.context_prompt
        current_model = model or state.model

    prompt = f"Meeting context: {ctx}\n\nTranscript: {transcript}\n\nRAG context: {rag_context}\n\nQuestion: {question}\n\nGive a concise, confident answer (under 80 words):"

    async def generate():
        try:
            response = completion(
                model_key=current_model,
                messages=[{"role": "user", "content": prompt}],
                stream=True, max_tokens=200,
            )
            for chunk in response:
                delta = chunk.choices[0].delta.content
                if delta:
                    yield f"data: {json.dumps({'text': delta})}\n\n"
            yield "data: {\"done\": true}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'error': str(e)})}\n\n"

    return StreamingResponse(generate(), media_type="text/event-stream",
                             headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})

@app.post("/meeting/summarize")
async def summarize_meeting(req: SummaryRequest):
    transcript_text = "\n".join(f"{t['speaker']}: {t['text']}" for t in req.transcript)
    if not HAS_LLM:
        return {"notes": _mock_summary(transcript_text)}
    try:
        prompt = SUMMARY_PROMPT.format(running_summary=state.running_summary, new_segment=transcript_text)
        response = completion(
            model_key=state.model,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=800,
        )
        notes = response.choices[0].message.content
        with state.lock:
            state.running_summary = notes
        return {"notes": notes}
    except Exception as e:
        return {"notes": _mock_summary(transcript_text), "error": str(e)}

def _mock_summary(transcript: str) -> str:
    return f"""## Meeting Notes — {time.strftime('%B %d, %Y')}

### Key Points
- Discussed system architecture and technical implementation
- Reviewed distributed tracing and observability approaches
- Covered authentication patterns for microservices

### Action Items
- [ ] Share architecture diagram by end of week
- [ ] Follow up on technical deep-dive scheduling

### Decisions
- API Gateway pattern selected for service mesh
- OpenTelemetry + Jaeger for distributed tracing"""


@app.get("/meeting/export")
async def export_notes(format: str = "md"):
    """Export meeting notes. format: md | pdf | docx"""
    notes = state.running_summary or _mock_summary("")
    export_dir = Path.home() / ".meetai" / "exports"
    export_dir.mkdir(parents=True, exist_ok=True)

    if format == "md":
        path = export_dir / f"meeting_{int(time.time())}.md"
        path.write_text(notes, encoding="utf-8")
        return FileResponse(path, media_type="text/markdown", filename=path.name)

    elif format == "pdf":
        try:
            from fpdf import FPDF
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Helvetica", size=14)
            pdf.set_text_color(30, 30, 30)
            for line in notes.split("\n"):
                clean = line.lstrip("# ").strip()
                if not clean:
                    pdf.ln(3)
                    continue
                if line.startswith("## "):
                    pdf.set_font("Helvetica", "B", 14); pdf.cell(0, 8, clean, ln=True)
                elif line.startswith("### "):
                    pdf.set_font("Helvetica", "B", 12); pdf.cell(0, 7, clean, ln=True)
                else:
                    pdf.set_font("Helvetica", "", 10); pdf.multi_cell(0, 5, clean)
            path = export_dir / f"meeting_{int(time.time())}.pdf"
            pdf.output(str(path))
            return FileResponse(path, media_type="application/pdf", filename=path.name)
        except ImportError:
            raise HTTPException(503, "fpdf2 not installed — pip install fpdf2")

    elif format == "docx":
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            doc = Document()
            for line in notes.split("\n"):
                if line.startswith("## "):
                    p = doc.add_heading(line[3:], 1); p.runs[0].font.color.rgb = RGBColor(0x63, 0x66, 0xF1)
                elif line.startswith("### "):
                    p = doc.add_heading(line[4:], 2)
                elif line.startswith("- [ ]"):
                    doc.add_paragraph(f"☐ {line[6:]}", style="List Bullet")
                elif line.startswith("- "):
                    doc.add_paragraph(f"• {line[2:]}", style="List Bullet")
                elif line.strip():
                    doc.add_paragraph(line)
            path = export_dir / f"meeting_{int(time.time())}.docx"
            doc.save(str(path))
            return FileResponse(path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=path.name)
        except ImportError:
            raise HTTPException(503, "python-docx not installed — pip install python-docx")

    raise HTTPException(400, f"Unsupported format: {format}. Use md, pdf, or docx.")


# ─── Screenshot / Vision Analysis ────────────────────────────────────────────

class ScreenshotRequest(BaseModel):
    image_b64: str
    path: Optional[str] = None
    prompt: str = "Describe what's on this screen. If there's a question, problem, or code, suggest the best response."

@app.post("/screenshot/analyze")
async def analyze_screenshot(req: ScreenshotRequest):
    """Accept base64 image and analyze with vision model."""
    if not HAS_LLM:
        return {"analysis": "⚠️ LLM not configured. Add ANTHROPIC_API_KEY or OPENAI_API_KEY to .env"}

    model_key = state.model

    # Ollama: only do vision if a vision model (e.g. llava) is configured
    if model_key == "ollama":
        _vision_model_name = os.environ.get("OLLAMA_VISION_MODEL", "")
        if not _vision_model_name:
            return {
                "analysis": "Set OLLAMA_VISION_MODEL=llava in .env to enable screenshot analysis. Run: ollama pull llava",
                "model": "none",
            }
        try:
            base = os.environ.get("OLLAMA_API_BASE", "http://localhost:11434")
            client = _openai_sdk.OpenAI(base_url=f"{base}/v1", api_key="ollama")
            messages = [{
                "role": "user",
                "content": [
                    {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{req.image_b64}"}},
                    {"type": "text", "text": req.prompt},
                ],
            }]
            resp = client.chat.completions.create(model=_vision_model_name, messages=messages, max_tokens=500)
            return {"analysis": resp.choices[0].message.content, "model": _vision_model_name}
        except Exception as exc:
            return {"analysis": f"Vision analysis failed: {exc}", "error": str(exc)}

    # Cloud models (claude / gpt4) — use standard vision messages
    try:
        messages = [{
            "role": "user",
            "content": [
                {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{req.image_b64}"}},
                {"type": "text", "text": req.prompt},
            ],
        }]
        response = completion(model_key=model_key, messages=messages, max_tokens=500)
        return {"analysis": response.choices[0].message.content, "model": MODEL_MAP.get(model_key, model_key)}
    except Exception as exc:
        return {"analysis": f"Vision analysis failed: {exc}", "error": str(exc)}


# ─── Action Items ─────────────────────────────────────────────────────────────

@app.get("/meeting/action-items")
async def get_action_items():
    """Return structured action items extracted from meeting notes."""
    from backend.summarizer.rolling import RollingSummarizer
    summarizer = RollingSummarizer(model=state.model)
    summarizer.running_summary = state.running_summary
    items = summarizer.extract_action_items()
    return {"action_items": items, "count": len(items)}


@app.post("/meeting/summarize/rolling")
async def rolling_summarize():
    """Force push current transcript into rolling summary."""
    from backend.summarizer.rolling import RollingSummarizer
    summarizer = RollingSummarizer(model=state.model)
    with state.lock:
        for entry in state.transcript:
            summarizer.add_segment(entry["speaker"], entry["text"])
    notes = summarizer.get_full_summary()
    with state.lock:
        state.running_summary = notes
    return {"notes": notes}


def _current_time() -> str:
    return time.strftime("%H:%M")


# ─── Entry point ─────────────────────────────────────────────────────────────

if __name__ == "__main__":
    uvicorn.run("backend.server:app", host="127.0.0.1", port=8765, reload=True, log_level="info")
